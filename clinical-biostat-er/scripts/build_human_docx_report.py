#!/usr/bin/env python3
from __future__ import annotations

import csv
import os
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[2]
BUNDLE = REPO / "clinical-biostat-er"
RUN = BUNDLE / "evals/_runs/pipeline_scaffold_mock01_review_20260620_182138"
LATEST = BUNDLE / "evals/visual_review/mock_dataset_01/comparison_packs/latest"
OUT_DIR = BUNDLE / "dist/mock01-human-readable-report_20260620_182138"
DOCX = OUT_DIR / "Mock01_项目说明与结果概览.docx"
TABLE_DOCX = OUT_DIR / "主结果_表格对照.docx"
EXAMPLE_DOCX = OUT_DIR / "报告示例_Origin对照.docx"
RUN_LABEL = "pipeline_scaffold_mock01_review_20260620_182138"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    for paragraph in cell.paragraphs:
      paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 17, "1F4D78", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 4),
        ("Heading 3", 11.5, "1F4D78", 6, 3),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mock01 项目说明与结果概览")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("面向第一次接触项目的 10 分钟阅读版 | 2026-06-20 最新 mock01 结果")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    add_callout(
        doc,
        "一句话说明",
        "这个项目把 AZ 原来的暴露-反应分析 Rmd，拆成一组可复用的分析和绘图 tools。"
        "Claude Code 不再临时写图，而是按固定入口读取数据、运行阶段化 workflow、生成表格/图、再用证据文件检查结果。"
    )


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_stage_table(doc: Document) -> None:
    stages = [
        ("1. 读入项目", "AZ mock 数据、Rmd、配置", "确认有哪些数据集、变量、endpoint、依赖"),
        ("2. 个体 PK/PD 预览", "个体浓度、给药、事件、response", "生成个体 profile 和 swimmer preview"),
        ("3. 暴露指标准备", "PK 暴露、剂量、事件时间", "形成可被后续模型复用的 subject-level exposure frame"),
        ("4. 暴露-反应探索", "暴露 frame + endpoint", "生成 ER 三联图：分组箱线、logistic、剂量分布"),
        ("5. 统计建模", "posthoc exposure data", "生成 logistic / Cox / KM 表格和 KM 图"),
        ("6. 复核打包", "Results tables/figures", "产出报告、图表、审计证据和 reviewer handoff"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["阶段", "输入", "产出/结果长什么样"]
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for stage, inp, out in stages:
        cells = table.add_row().cells
        for idx, text in enumerate([stage, inp, out]):
            set_cell_text(cells[idx], text)
    set_table_widths(table, [1.25, 2.0, 3.25])


def add_simple_table(doc: Document, rows: list[tuple[str, str]], widths=(1.8, 4.7)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], k, bold=True)
        set_cell_shading(cells[0], "F2F4F7")
        set_cell_text(cells[1], v)
    set_table_widths(table, list(widths))


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("555555")


def add_run_font(run, size: float = 9.0, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def add_csv_table(doc: Document, csv_path: Path, max_rows: int = 12,
                  max_cols: int = 7) -> None:
    rows = read_csv(csv_path)
    if not rows:
        doc.add_paragraph("（表格为空或未找到）")
        return
    cols = list(rows[0].keys())[:max_cols]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        set_cell_text(table.rows[0].cells[i], col, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            value = str(row.get(col, ""))
            if len(value) > 80:
                value = value[:77] + "..."
            set_cell_text(cells[i], value)
    if len(rows) > max_rows or len(rows[0].keys()) > max_cols:
        note = doc.add_paragraph(
            f"注：为保证 Word 可读性，此处展示前 {min(max_rows, len(rows))} 行、"
            f"前 {min(max_cols, len(rows[0].keys()))} 列；完整 CSV 仍在源运行目录中。"
        )
        note.runs[0].font.size = Pt(8.5)
        note.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_how_to_use(doc: Document) -> None:
    doc.add_heading("怎么用：人和 Claude Code 如何互动", level=1)
    add_simple_table(doc, [
        ("你给什么", "一个 AZ 项目目录：SourceData、Models、Scripts/ER_mock_analysis.Rmd，以及 workflow 配置。"),
        ("Claude Code 做什么", "按 skill 入口运行 staged workflow；它可以选择已有 tools，但不能临时粘贴新的 deliverable plotting code。"),
        ("系统输出什么", "Results/tables、Results/figures、intermediate 审计文件、comparison pack、人工报告。"),
        ("人需要看什么", "先看本 DOCX 的阶段解释和代表性结果；有争议时再打开 evidence appendix。"),
    ])
    doc.add_paragraph()
    add_callout(
        doc,
        "关键边界",
        "Claude Code 是 runner，不是绘图代码作者。真正的图由从 AZ Rmd 抽出的 builder-owned tools 生成；"
        "这就是我们这轮修正的重点。"
    )


def add_status_section(doc: Document) -> None:
    table_rows = read_csv(LATEST / "results_table_diff_summary.csv")
    figure_rows = read_csv(LATEST / "figure_input_accuracy_summary.csv")
    pass_rows = sum(1 for r in figure_rows if r.get("primary_issue_class") == "pass_current_boundary")
    review_rows = len(figure_rows) - pass_rows
    doc.add_heading("现在做到什么程度", level=1)
    add_simple_table(doc, [
        ("表格结果", f"{len(table_rows)} 张 Results tables 已按 AZ reference 对齐。这里的表格包括 logistic、Cox、KM、ER summary 等统计结果。"),
        ("图片生成", f"{len(figure_rows)} 张图都有对应输出；当前全部使用从 AZ Rmd 抽出的 direct plotting tools。"),
        ("当前边界", f"{pass_rows} 张在当前自动边界内通过；{review_rows} 张 Core2 preview 仍需临床语义/adapter review。"),
        ("不能宣称", "这不是最终临床决策包，也不是逐图 layer-level pixel parity 证明。它是 review-ready 的复刻工作流演示。"),
    ])


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("这到底是什么", level=1)
    doc.add_paragraph(
        "这是一个临床药理/生物统计暴露-反应分析的自动化复刻工作流。"
        "我们不是重新发明统计分析，而是把 AZ 已有 Rmd 里的数据准备、统计表格、绘图逻辑拆成稳定 tools，"
        "让 Claude Code 可以按固定流程运行、产出结果，并留下可检查的证据。"
    )
    add_bullets(doc, [
        "适合的场景：已有 reference Rmd，希望复刻、迁移、审查、或让 agent 稳定重复运行。",
        "不适合的场景：让 agent 自己发挥、临时写全新统计图、直接跳到临床结论。",
        "当前 mock01 用途：证明这条 workflow 能从输入数据跑到 Results 表格/图片，并能解释每一步的证据边界。",
    ])

    add_how_to_use(doc)

    doc.add_heading("输入和输出是什么", level=1)
    add_simple_table(doc, [
        ("主要输入", "AZ mock 数据目录、原始 ER_mock_analysis.Rmd、模型/posthoc 暴露数据、workflow 配置。"),
        ("中间产物", "数据盘点、暴露指标、个体图输入、posthoc exposure frame、模型诊断、方法选择审计。"),
        ("最终产物", "Results/tables 下的统计表格，Results/figures 下的图，以及给 reviewer 的报告包。"),
        ("评估证据", "comparison pack、figure input audit、plot capability ownership map、direct extract backlog。"),
    ])

    doc.add_heading("分几个阶段，每个阶段长什么样", level=1)
    add_stage_table(doc)

    doc.add_heading("阶段结果示例", level=1)
    doc.add_heading("阶段 2：个体 PK/PD 和事件预览", level=2)
    doc.add_paragraph(
        "这一阶段帮助 reviewer 看单个受试者的给药、浓度、response、AE 时间线是否被正确整理。"
        "图本身现在调用 AZ Rmd 中抽出的 profile/swimmer plotting functions。"
    )
    add_figure(
        doc,
        LATEST / "20250925_pkind4__pipeline_scaffold_mock01_review_20260620_182138.png",
        "示例：个体 PK profile preview。用于检查个体浓度、时间点和事件标注。"
    )
    add_figure(
        doc,
        LATEST / "swimmer_high_dose__pipeline_scaffold_mock01_review_20260620_182138.png",
        "示例：high-dose swimmer preview。用于查看治疗持续时间、response 和事件 overlay。"
    )

    doc.add_heading("阶段 4：暴露-反应三联图", level=2)
    doc.add_paragraph(
        "这一阶段把暴露指标和 endpoint 连起来。典型图包括左侧暴露分布比较、右上 logistic exposure-response 曲线、右下剂量组分布。"
        "Core4 的 32 张 ER 图已切换为 AZ Rmd direct-extracted plotting tool。"
    )
    add_figure(
        doc,
        LATEST / "ER_AUC1_Res1_efficacy__pipeline_scaffold_mock01_review_20260620_182138.png",
        "示例：AUC1 vs confirmed response 的 ER 三联图。"
    )

    doc.add_heading("阶段 5：KM / Cox / TTE 图和统计表", level=2)
    doc.add_paragraph(
        "这一阶段生成 OS、PFS、DoR、ILD 等时间到事件相关图和表。"
        "Core5 的 KM/TTE 图现在也走 AZ Rmd 中的 survfit / ggsurvplot / ggarrange direct tools。"
    )
    add_figure(
        doc,
        LATEST / "Combined_ILD_incidence_curves__pipeline_scaffold_mock01_review_20260620_182138.png",
        "示例：ILD cumulative incidence combined plot。"
    )

    add_status_section(doc)

    doc.add_heading("我们怎么评估", level=1)
    doc.add_paragraph(
        "评估分两层。第一层是机器可重复检查：文件有没有生成、输入 frame 是否存在、必要列是否齐全、表格数值是否和 reference 对齐、图是不是来自 AZ direct tools。"
        "第二层是人工 review：某个图的临床语义、图层映射和最终解释是否合理。两层不能混为一谈。"
    )
    add_simple_table(doc, [
        ("机器检查回答的问题", "流程有没有跑通？结果文件有没有生成？表格数值是否复刻？图是否调用正确 tool？"),
        ("人工 review 回答的问题", "图是否表达了正确临床含义？某个 endpoint、event window、adapter 是否符合业务理解？"),
        ("当前结论", "适合给 reviewer 看和继续检查；不直接作为最终医学/注册决策材料。"),
    ])

    doc.add_heading("这次最重要的修正", level=1)
    add_callout(
        doc,
        "从“agent 会画图”改成“AZ 绘图代码变成 tools”",
        "之前的问题是把“图生成了”误当成“图的绘图逻辑正确”。现在 mock01 的 54 张图都标记为 az_rmd_direct："
        "Core2、Core4、Core5 都不再依赖 runner 临时写绘图代码，而是调用从 AZ Rmd 抽出的稳定 plotting tools。"
    )

    doc.add_heading("给第一次接触项目的人怎么读", level=1)
    add_bullets(doc, [
        "先读前两页：理解它是什么、输入输出是什么、人和 Claude Code 怎么互动。",
        "再看阶段表和 4 张示例图：理解每个阶段产物长什么样。",
        "最后看“怎么评估”和“当前边界”：知道哪些已经自动证明，哪些还需要人工判断。",
        "只有当你要追证据时，才打开 comparison pack、CSV audit 或原始 Results 文件。"
    ])

    doc.add_heading("附录：最少必要证据位置", level=1)
    add_simple_table(doc, [
        ("主结果表格对照", "主结果_表格对照.docx：9 张 Results tables，全部按 origin/generated 1 对 1 成组保存。"),
        ("报告示例对照", "报告示例_Origin对照.docx：每个图类别一个 origin vs generated 例子，另加一组表格对照示例。"),
        ("主运行目录", str(RUN)),
        ("最新 comparison appendix", str(LATEST / "index.html")),
        ("figure input audit", str(LATEST / "figure_input_accuracy_summary.csv")),
        ("table diff summary", str(LATEST / "results_table_diff_summary.csv")),
        ("direct extract coverage", str(BUNDLE / "evals/_runs/mock01_review_acceptance_20260620_182138/plot_capability_extraction_coverage.csv")),
    ], widths=(1.65, 4.85))

    doc.save(DOCX)


def build_package() -> None:
    build_table_comparison_docx()
    build_example_comparison_docx()
    write_package_readme()

    package_zip = OUT_DIR.with_suffix(".zip")
    if package_zip.exists():
        package_zip.unlink()
    with zipfile.ZipFile(package_zip, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for path in [DOCX, TABLE_DOCX, EXAMPLE_DOCX, OUT_DIR / "README_先读这个.md"]:
            if path.exists():
                z.write(path, path.name)


def table_pair_names() -> list[str]:
    return [
        "Cox_PH_models_PFS_OS_summary.csv",
        "Enhanced_ER_analysis_summary.csv",
        "Final_Logistic_Regression_Complete_Results.csv",
        "Final_Logistic_Regression_Detailed_Summary.csv",
        "Final_Logistic_Regression_P_Values_Summary.csv",
        "ILD_Cox_regression_results.csv",
        "ILD_KM_analysis_summary.csv",
        "KM_analysis_summary_by_dose_stratification.csv",
        "KM_analysis_summary_Cave0_and_AUC1_twotiles_with_DoR.csv",
    ]


def build_table_comparison_docx() -> None:
    doc = Document()
    style_doc(doc)
    p = doc.add_paragraph()
    r = p.add_run("主结果：Results Tables 的 Origin / Generated 对照")
    add_run_font(r, 20, True)
    doc.add_paragraph(
        "这个文档把 9 张 Results tables 按 origin/generated 成组放在一起。"
        "表格内容是可编辑的 Word 表格；为了保持文档可读，较大的表只展示前几行/列。"
    )
    table_pairs = [
        *table_pair_names()
    ]
    for idx, name in enumerate(table_pairs, start=1):
        stem = Path(name).stem
        doc.add_heading(f"{idx}. {name}", level=1)
        doc.add_heading("Origin / AZ reference", level=2)
        add_csv_table(doc, LATEST / f"{stem}__original.csv")
        doc.add_heading("Generated / workflow output", level=2)
        add_csv_table(doc, LATEST / f"{stem}__{RUN_LABEL}.csv")
        if idx != len(table_pairs):
            doc.add_page_break()
    doc.save(TABLE_DOCX)


def add_origin_generated_figure_pair(doc: Document, title: str, file_name: str) -> None:
    stem = Path(file_name).stem
    ext = Path(file_name).suffix
    doc.add_heading(title, level=1)
    doc.add_paragraph("左侧是 origin/AZ reference，右侧是本轮 workflow generated。图片嵌入在 Word 中，说明文字可直接编辑。")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["Origin", "Generated"]):
        set_cell_text(table.rows[0].cells[i], label, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    paths = [
        LATEST / f"{stem}__original{ext}",
        LATEST / f"{stem}__{RUN_LABEL}{ext}",
    ]
    for i, path in enumerate(paths):
        cell = table.rows[1].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if path.exists():
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Inches(3.0))
        else:
            set_cell_text(cell, f"未找到：{path.name}")
    doc.add_paragraph()


def build_example_comparison_docx() -> None:
    doc = Document()
    style_doc(doc)
    p = doc.add_paragraph()
    r = p.add_run("报告示例：每类结果的 Origin / Generated 对照")
    add_run_font(r, 20, True)
    doc.add_paragraph(
        "这个文档只放少量示例，用来解释不同阶段的结果长什么样。"
        "它不是完整证据包；完整 9 张表格对照在《主结果_表格对照.docx》。"
    )
    examples = [
        ("Core2 个体 PK 曲线示例", "20250925_pkind4.png"),
        ("Core2 swimmer 事件图示例", "swimmer_high_dose.png"),
        ("Core4 ER 三联图示例", "ER_AUC1_Res1_efficacy.png"),
        ("Core5 KM 生存曲线示例", "OS_KM_Cave_0_to_OS_twotiles.png"),
        ("Core5 ILD 累计发生率示例", "Combined_ILD_incidence_curves.png"),
    ]
    for idx, (title, file_name) in enumerate(examples, start=1):
        add_origin_generated_figure_pair(doc, title, file_name)
        if idx in (2, 4):
            doc.add_page_break()

    doc.add_heading("表格对照示例", level=1)
    example_table = "Final_Logistic_Regression_Complete_Results.csv"
    example_stem = Path(example_table).stem
    doc.add_heading("Origin / AZ reference", level=2)
    add_csv_table(doc, LATEST / f"{example_stem}__original.csv", max_rows=10, max_cols=6)
    doc.add_heading("Generated / workflow output", level=2)
    add_csv_table(doc, LATEST / f"{example_stem}__{RUN_LABEL}.csv", max_rows=10, max_cols=6)
    doc.save(EXAMPLE_DOCX)


def write_package_readme() -> None:
    readme = [
        "# 报告包目录说明",
        "",
        "这个包只保留可编辑的 Word 文档，不直接交付散落的 PNG 文件。",
        "",
        "## 文件",
        "- `Mock01_项目说明与结果概览.docx`：给第一次接触项目的人看的中文主报告。",
        "- `主结果_表格对照.docx`：9 张 Results tables 的 origin/generated 成组对照，表格为 Word 可编辑表格。",
        "- `报告示例_Origin对照.docx`：每类图一个 origin/generated 示例，另加一组表格对照示例。",
        "",
        "注：图在 Word 中作为图片嵌入，用于报告展示；文字、标题、说明和表格内容都可以在 Word 中编辑。"
    ]
    (OUT_DIR / "README_先读这个.md").write_text("\n".join(readme), encoding="utf-8")


if __name__ == "__main__":
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    build_docx()
    build_package()
    print(f"Wrote {DOCX}")
    print(f"Wrote {OUT_DIR.with_suffix('.zip')}")
