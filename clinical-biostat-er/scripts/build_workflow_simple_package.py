#!/usr/bin/env python3
from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BUNDLE = Path(__file__).resolve().parents[1]
RUN_LABEL = "pipeline_scaffold_mock01_review_20260620_182138"
RUN = BUNDLE / "evals/_runs" / RUN_LABEL
LATEST = BUNDLE / "evals/visual_review/mock_dataset_01/comparison_packs/latest"
HUMAN_REPORT = BUNDLE / "dist/mock01-human-readable-report_20260620_182138"
OUT_DIR = BUNDLE / "dist/mock01-workflow-simple-package_20260620_182138"
DOCX = OUT_DIR / "Mock01_Workflow_简化结果说明.docx"
ZIP = OUT_DIR.with_suffix(".zip")


STEPS = [
    ("01", "理解数据", "确认数据集、变量、endpoint、暴露指标和质量问题。", [
        RUN / "intermediate/01_understanding_data/dataset_inventory.csv",
        RUN / "intermediate/01_understanding_data/endpoint_inventory.csv",
        RUN / "intermediate/01_understanding_data/data_quality_findings.csv",
    ]),
    ("02", "个体 PK/PD 预览", "生成 Core2 profile/swimmer 代表性预览。", [
        LATEST / "20250925_pkind4__original.png",
        LATEST / f"20250925_pkind4__{RUN_LABEL}.png",
        LATEST / "swimmer_high_dose__original.png",
        LATEST / f"swimmer_high_dose__{RUN_LABEL}.png",
    ]),
    ("03", "暴露指标", "形成后续 ER/modeling 可复用的 subject-level exposure frame。", [
        RUN / "intermediate/03_exposure_metrics/exposure_metric_definitions.csv",
        RUN / "intermediate/03_exposure_metrics/subject_exposure_metrics.csv",
    ]),
    ("04", "暴露-反应探索", "生成 ER pair 图和 ER summary。", [
        LATEST / "ER_AUC1_Res1_efficacy__original.png",
        LATEST / f"ER_AUC1_Res1_efficacy__{RUN_LABEL}.png",
        LATEST / "Enhanced_ER_analysis_summary__original.csv",
        LATEST / f"Enhanced_ER_analysis_summary__{RUN_LABEL}.csv",
    ]),
    ("05", "统计建模", "生成 logistic/Cox/KM 结果表和 KM 图。", [
        LATEST / "PFS_KM_by_dose__original.png",
        LATEST / f"PFS_KM_by_dose__{RUN_LABEL}.png",
        LATEST / "Final_Logistic_Regression_Complete_Results__original.csv",
        LATEST / f"Final_Logistic_Regression_Complete_Results__{RUN_LABEL}.csv",
    ]),
    ("06", "复核与报告", "生成 review evidence 和最终中文 DOCX 报告。", [
        LATEST / "results_table_diff_summary.csv",
        LATEST / "figure_input_accuracy_summary.csv",
        HUMAN_REPORT / "Mock01_项目说明与结果概览.docx",
    ]),
]


TABLE_BASES = [
    "Cox_PH_models_PFS_OS_summary",
    "Enhanced_ER_analysis_summary",
    "Final_Logistic_Regression_Complete_Results",
    "Final_Logistic_Regression_Detailed_Summary",
    "Final_Logistic_Regression_P_Values_Summary",
    "ILD_Cox_regression_results",
    "ILD_KM_analysis_summary",
    "KM_analysis_summary_Cave0_and_AUC1_twotiles_with_DoR",
    "KM_analysis_summary_by_dose_stratification",
]


def safe_name(text: str) -> str:
    return (
        text.replace("/", "")
        .replace("\\", "")
        .replace(" ", "")
        .replace(":", "")
    )


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def copy_file(src: Path, dst: Path) -> bool:
    if not src.exists() or not src.is_file():
        return False
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)
    return True


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def set_run_font(run, size: float = 10, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.6) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [("Heading 1", 16, "0B2545"), ("Heading 2", 12.5, "1F4D78")]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FB")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 10.2, True, "0B2545")
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, 9.6)
    doc.add_paragraph()


def add_csv_preview(doc: Document, path: Path, rows_n: int = 4, cols_n: int = 5) -> None:
    if not path.exists() or path.suffix.lower() != ".csv":
        return
    rows = read_csv(path)
    if not rows:
        return
    cols = list(rows[0].keys())[:cols_n]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        set_cell_text(table.rows[0].cells[i], col, True, 7.8)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows[:rows_n]:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            value = str(row.get(col, ""))
            if len(value) > 45:
                value = value[:42] + "..."
            set_cell_text(cells[i], value, False, 7.3)
    p = doc.add_paragraph(f"预览：{path.name}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, 8.2, False, "666666")


def add_image_pair(doc: Document, origin: Path, ours: Path, caption: str) -> None:
    if not origin.exists() or not ours.exists():
        return
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(("Origin", "我们的输出")):
        set_cell_text(table.rows[0].cells[i], h, True, 8.5)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for i, path in enumerate((origin, ours)):
        cell = table.rows[1].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.9))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, 8.4, False, "666666")


def build_tree() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    ensure_dir(OUT_DIR)

    for step_id, title, purpose, files in STEPS:
        step_dir = OUT_DIR / "examples_by_step" / f"{step_id}_{safe_name(title)}"
        ensure_dir(step_dir)
        lines = [f"# {step_id} {title}", "", purpose, ""]
        for src in files:
            prefix = "origin__" if "__original" in src.name else "ours__"
            if src.name in {"results_table_diff_summary.csv", "figure_input_accuracy_summary.csv"}:
                prefix = "evidence__"
            if "Mock01_" in src.name:
                prefix = "final_report__"
            dst = step_dir / f"{prefix}{src.name}"
            if copy_file(src, dst):
                lines.append(f"- `{dst.name}`")
        (step_dir / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    tables_dir = OUT_DIR / "all_result_tables_pairs"
    for base in TABLE_BASES:
        pair_dir = tables_dir / base
        copy_file(LATEST / f"{base}__original.csv", pair_dir / f"origin__{base}.csv")
        copy_file(LATEST / f"{base}__{RUN_LABEL}.csv", pair_dir / f"ours__{base}.csv")

    final_dir = OUT_DIR / "final_report_docx"
    for src in sorted(HUMAN_REPORT.glob("*.docx")):
        copy_file(src, final_dir / src.name)

    source_dir = OUT_DIR / "workflow_source"
    copy_file(RUN / "analysis/er_core_workflow.Rmd", source_dir / "er_core_workflow.Rmd")
    copy_file(RUN / "config/er_workflow_spec.yaml", source_dir / "er_workflow_spec.yaml")
    for rel in (
        "skills/er-individual-pk-pd-review/code_corpus/az_mock01_core2_reference_plotters.R",
        "skills/er-exposure-response-exploration/code_corpus/az_mock01_core4_er_plotters.R",
        "skills/er-statistical-modeling/code_corpus/az_mock01_core5_km_plotters.R",
    ):
        copy_file(BUNDLE / rel, source_dir / Path(rel).name)

    (OUT_DIR / "README_先读这个.md").write_text(
        "# Mock01 workflow 简化结果包\n\n"
        "先打开 `Mock01_Workflow_简化结果说明.docx`。\n\n"
        "包里只保留四个入口：\n\n"
        "- `examples_by_step/`: 每个 workflow 步骤一个小文件夹，只放代表性结果。\n"
        "- `all_result_tables_pairs/`: 9 张 Results tables 的 origin/ours 1:1 对照。\n"
        "- `final_report_docx/`: 最终中文报告 DOCX。\n"
        "- `workflow_source/`: workflow Rmd、配置和从 AZ 抽出的绘图 tools。\n",
        encoding="utf-8",
    )


def build_docx() -> None:
    doc = Document()
    style_doc(doc)
    p = doc.add_paragraph()
    r = p.add_run("Mock01 Workflow 简化结果说明")
    set_run_font(r, 22, True, "0B2545")
    p = doc.add_paragraph()
    r = p.add_run(f"运行结果：{RUN_LABEL} | 这个版本把复杂 evidence 包简化成 4 个入口")
    set_run_font(r, 10.5, False, "555555")
    add_callout(
        doc,
        "一句话",
        "这个包不是让读者审计所有 CSV，而是让读者按 workflow 顺序快速看到：每一步做了什么、代表性结果长什么样、"
        "哪些地方有 origin 对照、最终中文报告在哪里。"
    )

    doc.add_heading("包里只有四个入口", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(("入口", "里面是什么", "什么时候看")):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "DCE9F6")
    rows = [
        ("examples_by_step", "6 个步骤的代表性结果；没有再套 origin/our/evidence 多层目录。", "先看这里。"),
        ("all_result_tables_pairs", "9 张 Results tables，每张表一个文件夹，里面只有 origin 和 ours 两个 CSV。", "需要核对表格结果时看。"),
        ("final_report_docx", "最终中文 DOCX 报告。", "给别人展示或继续编辑时看。"),
        ("workflow_source", "workflow Rmd、配置、AZ plotter tools。", "需要追代码来源时看。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    doc.add_heading("Workflow 六步", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("步骤", "做什么", "代表性文件夹", "origin 情况")):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "DCE9F6")
    for step_id, title, purpose, files in STEPS:
        has_origin = any("__original" in p.name for p in files)
        cells = table.add_row().cells
        for i, text in enumerate((f"{step_id} {title}", purpose, f"examples_by_step/{step_id}_{safe_name(title)}", "有代表性 origin 对照" if has_origin else "没有逐文件 origin")):
            set_cell_text(cells[i], text, False, 8.3)

    doc.add_heading("代表性图结果", level=1)
    add_image_pair(
        doc,
        LATEST / "20250925_pkind4__original.png",
        LATEST / f"20250925_pkind4__{RUN_LABEL}.png",
        "Step 2: Core2 individual PK profile preview",
    )
    add_image_pair(
        doc,
        LATEST / "ER_AUC1_Res1_efficacy__original.png",
        LATEST / f"ER_AUC1_Res1_efficacy__{RUN_LABEL}.png",
        "Step 4: ER pair figure",
    )
    add_image_pair(
        doc,
        LATEST / "PFS_KM_by_dose__original.png",
        LATEST / f"PFS_KM_by_dose__{RUN_LABEL}.png",
        "Step 5: KM figure",
    )

    doc.add_heading("代表性表格结果", level=1)
    doc.add_paragraph("Origin: Final_Logistic_Regression_Complete_Results")
    add_csv_preview(doc, LATEST / "Final_Logistic_Regression_Complete_Results__original.csv")
    doc.add_paragraph("我们的输出: Final_Logistic_Regression_Complete_Results")
    add_csv_preview(doc, LATEST / f"Final_Logistic_Regression_Complete_Results__{RUN_LABEL}.csv")

    add_callout(
        doc,
        "最终 report",
        "`final_report_docx/Mock01_项目说明与结果概览.docx` 是最终 workflow 的人类阅读入口；"
        "`主结果_表格对照.docx` 和 `报告示例_Origin对照.docx` 是它的两个配套 DOCX。"
    )
    doc.save(DOCX)


def zip_package() -> None:
    if ZIP.exists():
        ZIP.unlink()
    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.rglob("*")):
            if path.is_file() and not any(part.startswith("render_") for part in path.parts):
                zf.write(path, path.relative_to(OUT_DIR.parent))


def main() -> None:
    build_tree()
    build_docx()
    zip_package()
    print(OUT_DIR)
    print(ZIP)


if __name__ == "__main__":
    main()
