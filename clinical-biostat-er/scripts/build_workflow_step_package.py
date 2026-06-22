#!/usr/bin/env python3
from __future__ import annotations

import csv
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BUNDLE = Path(__file__).resolve().parents[1]
RUN_LABEL = "pipeline_scaffold_mock01_review_20260620_182138"
RUN = BUNDLE / "evals/_runs" / RUN_LABEL
LATEST = BUNDLE / "evals/visual_review/mock_dataset_01/comparison_packs/latest"
HUMAN_REPORT = BUNDLE / "dist/mock01-human-readable-report_20260620_182138"
OUT_DIR = BUNDLE / "dist/mock01-workflow-step-results-package_20260620_182138"
DOCX = OUT_DIR / "Workflow_分步骤结果对照说明.docx"
ZIP = OUT_DIR.with_suffix(".zip")


@dataclass(frozen=True)
class Step:
    key: str
    title: str
    purpose: str
    inputs: str
    outputs: str
    origin_note: str
    generated_dirs: tuple[str, ...]
    origin_patterns: tuple[str, ...] = ()
    comparison_patterns: tuple[str, ...] = ()
    evidence_files: tuple[str, ...] = ()
    representative_files: tuple[str, ...] = ()


STEPS: tuple[Step, ...] = (
    Step(
        key="01_understanding_data",
        title="Step 1 - 理解数据与定义分析边界",
        purpose="先确认有哪些数据集、变量、endpoint、剂量、暴露指标、缺失和连接键问题。",
        inputs="项目配置、源数据入口、AZ Rmd/脚本依赖。",
        outputs="dataset/endpoint/exposure inventories、data quality findings、readiness flags。",
        origin_note="这一阶段没有逐文件 origin 输出对照；origin 侧主要是原始项目输入和 Rmd 语义。",
        generated_dirs=("intermediate/01_understanding_data",),
        evidence_files=("pipeline_status.csv",),
        representative_files=(
            "intermediate/01_understanding_data/dataset_inventory.csv",
            "intermediate/01_understanding_data/endpoint_inventory.csv",
            "intermediate/01_understanding_data/data_quality_findings.csv",
        ),
    ),
    Step(
        key="02_individual_pk_pd_review",
        title="Step 2 - 个体 PK/PD 与 Core2 reference preview",
        purpose="把个体给药、PK 点、response 和安全事件组织成 profile/swimmer 类预览图。",
        inputs="subject index、给药区间、PK 采样点、response/safety 事件。",
        outputs="Core2 profile/swimmer preview 图，以及对应 plotting calls、QC、manifest。",
        origin_note="origin 侧有 AZ reference preview 图；中间 frame 没有完整 origin 对照。",
        generated_dirs=("intermediate/02_individual_pk_pd_review", "outputs/02_individual_pk_pd_review"),
        origin_patterns=("20250925_pkind*", "pkind_payload_*", "swimmer_*"),
        comparison_patterns=("20250925_pkind*", "pkind_payload_*", "swimmer_*"),
        evidence_files=("figure_input_accuracy_summary.csv", "figure_semantic_contract.csv"),
        representative_files=(
            "intermediate/02_individual_pk_pd_review/individual_profile_preview_manifest.csv",
            "intermediate/02_individual_pk_pd_review/reference_figure_preview_manifest.csv",
        ),
    ),
    Step(
        key="03_exposure_metrics",
        title="Step 3 - 生成暴露指标",
        purpose="把 PK/剂量信息整理成后续 ER 和统计模型可复用的 subject-level exposure frame。",
        inputs="剂量、PK、subject-level 记录和暴露定义。",
        outputs="exposure metric definitions、records、subject_exposure_metrics。",
        origin_note="这一阶段没有逐文件 origin 输出对照；它是后续 ER/modeling 的输入准备层。",
        generated_dirs=("intermediate/03_exposure_metrics",),
        representative_files=(
            "intermediate/03_exposure_metrics/exposure_metric_definitions.csv",
            "intermediate/03_exposure_metrics/subject_exposure_metrics.csv",
        ),
    ),
    Step(
        key="04_exposure_response_exploration",
        title="Step 4 - 暴露-反应探索图与 ER summary",
        purpose="生成 AZ Rmd 对应的 ER pair 图：暴露分布、logistic/endpoint 曲线、剂量分布等。",
        inputs="exposure_for_join、response/safety endpoint、ER question matrix。",
        outputs="ER pair figures、ER question/method audit、Enhanced ER summary table。",
        origin_note="origin 侧有对应 ER 图和 Enhanced_ER_analysis_summary.csv。",
        generated_dirs=("intermediate/04_exposure_response_exploration", "Results/figures", "Results/tables"),
        origin_patterns=("ER_*", "Enhanced_ER_analysis_summary*"),
        comparison_patterns=("ER_*", "Enhanced_ER_analysis_summary*"),
        evidence_files=("results_table_diff_summary.csv", "figure_input_accuracy_summary.csv"),
        representative_files=(
            "intermediate/04_exposure_response_exploration/mock01_er_pair_figure_manifest.csv",
            "Results/tables/Enhanced_ER_analysis_summary.csv",
        ),
    ),
    Step(
        key="05_statistical_modeling",
        title="Step 5 - Logistic / Cox / KM 统计建模",
        purpose="生成 Results tables 和 KM/Cox/TTE 图，是统计结果复刻证据最强的一层。",
        inputs="posthoc exposure data、endpoint/censoring、dose/exposure stratification。",
        outputs="9 张 Results tables、KM/Cox/ILD figures、model diagnostics 和 manifests。",
        origin_note="origin 侧有 9 张结果表和 KM/Cox/ILD 图；这些和我们的输出一一成组保存。",
        generated_dirs=("intermediate/05_statistical_modeling", "Results/tables", "Results/figures", "outputs/05_statistical_modeling"),
        origin_patterns=(
            "Final_Logistic_*",
            "Cox_PH_*",
            "ILD_*",
            "KM_analysis_*",
            "PFS_*",
            "OS_*",
            "DoR_*",
            "Combined_*",
        ),
        comparison_patterns=(
            "Final_Logistic_*",
            "Cox_PH_*",
            "ILD_*",
            "KM_analysis_*",
            "PFS_*",
            "OS_*",
            "DoR_*",
            "Combined_*",
        ),
        evidence_files=("results_table_diff_summary.csv", "results_table_reproduction_readiness.csv", "figure_input_accuracy_summary.csv"),
        representative_files=(
            "intermediate/05_statistical_modeling/model_run_summary.csv",
            "intermediate/05_statistical_modeling/mock01_results_table_manifest.csv",
            "Results/tables/Final_Logistic_Regression_Complete_Results.csv",
            "Results/tables/Cox_PH_models_PFS_OS_summary.csv",
        ),
    ),
    Step(
        key="06_reporting_review",
        title="Step 6 - 复核、证据包与人工报告",
        purpose="把最终图表、审计证据、review gate 和面向人的报告组织起来。",
        inputs="Results tables/figures、comparison pack、review gate evidence。",
        outputs="review summary、comparison evidence、最终中文 DOCX report package。",
        origin_note="origin 侧主要是 comparison pack 的 reference/original outputs；最终中文报告是我们的 workflow 产物。",
        generated_dirs=("intermediate/06_reporting_review", "outputs/06_reporting_review"),
        origin_patterns=("*.csv", "*.png", "*.pdf"),
        comparison_patterns=("*",),
        evidence_files=(
            "manifest.csv",
            "coverage_summary.csv",
            "results_table_diff_summary.csv",
            "figure_input_accuracy_summary.csv",
            "figure_semantic_contract.csv",
            "figure_plotted_data_summary.csv",
        ),
        representative_files=(
            "intermediate/06_reporting_review/review_gate_summary.csv",
            "intermediate/06_reporting_review/deliverable_readiness.csv",
            "outputs/06_reporting_review/review_summary.md",
        ),
    ),
)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def copy_file(src: Path, dst: Path) -> None:
    if not src.exists() or not src.is_file():
        return
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)


def copy_tree_files(src_dir: Path, dst_dir: Path) -> int:
    if not src_dir.exists():
        return 0
    count = 0
    for src in sorted(p for p in src_dir.rglob("*") if p.is_file()):
        rel = src.relative_to(src_dir)
        copy_file(src, dst_dir / rel)
        count += 1
    return count


def copy_latest_patterns(patterns: tuple[str, ...], dst: Path, require_origin: bool | None = None) -> int:
    count = 0
    for pattern in patterns:
        for src in sorted(LATEST.glob(pattern)):
            if not src.is_file():
                continue
            name = src.name
            if require_origin is True and "__original" not in name:
                continue
            if require_origin is False and "__original" in name:
                continue
            copy_file(src, dst / name)
            count += 1
    return count


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text, encoding="utf-8")


def set_run_font(run, size: float = 10.0, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color, before, after in [
        ("Heading 1", 16, "0B2545", 14, 5),
        ("Heading 2", 12.5, "1F4D78", 9, 3),
        ("Heading 3", 11, "1F4D78", 5, 2),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Mock01 Workflow 分步骤结果对照包")
    set_run_font(r, 22, True, "0B2545")
    p = doc.add_paragraph()
    r = p.add_run(f"最新运行：{RUN_LABEL} | 目标：按 workflow 看每一步输入、输出、origin 对照和最终报告")
    set_run_font(r, 10.5, False, "555555")
    add_callout(
        doc,
        "怎么读这个包",
        "先读本 DOCX。然后按 steps/01 到 steps/06 打开文件夹。每一步都尽量分成 origin_available、our_outputs、comparison_examples 和 evidence；"
        "没有 origin 的地方会明确写 origin_not_available.md，不把结构检查伪装成结果对齐。"
    )


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FB")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 10.3, True, "0B2545")
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, 9.5)
    doc.add_paragraph()


def add_step_overview_table(doc: Document, step_counts: dict[str, dict[str, int]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["步骤", "阶段目的", "输入", "我们的输出", "origin 可用性", "文件数"]
    widths = [1.05, 1.65, 1.65, 1.95, 1.45, 0.75]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, 8.8)
        set_cell_shading(table.rows[0].cells[i], "DCE9F6")
        table.rows[0].cells[i].width = Inches(widths[i])
    for step in STEPS:
        counts = step_counts[step.key]
        cells = table.add_row().cells
        values = [
            step.key[:2],
            step.purpose,
            step.inputs,
            step.outputs,
            "有 origin 对照" if counts["origin"] else "无逐文件 origin",
            str(counts["total"]),
        ]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, False, 8.1)
            cells[i].width = Inches(widths[i])


def add_csv_preview(doc: Document, path: Path, max_rows: int = 5, max_cols: int = 5) -> None:
    if not path.exists() or path.suffix.lower() != ".csv":
        return
    rows = read_csv(path)
    if not rows:
        return
    cols = list(rows[0].keys())[:max_cols]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        set_cell_text(table.rows[0].cells[i], col, True, 7.8)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            value = str(row.get(col, ""))
            if len(value) > 48:
                value = value[:45] + "..."
            set_cell_text(cells[i], value, False, 7.4)
    note = doc.add_paragraph(f"预览：{path.name}（仅展示前 {min(max_rows, len(rows))} 行；完整文件在包内）。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        set_run_font(run, 8.2, False, "666666")


def add_image_pair(doc: Document, origin: Path, generated: Path, caption: str) -> None:
    if not origin.exists() or not generated.exists():
        return
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for i, header in enumerate(("Origin", "我们的输出")):
        set_cell_text(table.rows[0].cells[i], header, True, 8.5)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for idx, path in enumerate((origin, generated)):
        cell = table.rows[1].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.9))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, 8.5, False, "666666")


def build_package_tree() -> dict[str, dict[str, int]]:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    ensure_dir(OUT_DIR)
    step_counts: dict[str, dict[str, int]] = {}

    for step in STEPS:
        root = OUT_DIR / "steps" / step.key
        origin_dir = root / "origin_available"
        our_dir = root / "our_outputs"
        comp_dir = root / "comparison_examples"
        evidence_dir = root / "evidence"
        ensure_dir(origin_dir)
        ensure_dir(our_dir)
        ensure_dir(comp_dir)
        ensure_dir(evidence_dir)

        generated_count = 0
        for rel_dir in step.generated_dirs:
            generated_count += copy_tree_files(RUN / rel_dir, our_dir / rel_dir.replace("/", "__"))

        origin_count = copy_latest_patterns(step.origin_patterns, origin_dir, require_origin=True)
        comp_count = copy_latest_patterns(step.comparison_patterns, comp_dir, require_origin=None)
        evidence_count = 0
        for name in step.evidence_files:
            src = LATEST / name
            if src.exists():
                copy_file(src, evidence_dir / name)
                evidence_count += 1
            src = RUN / name
            if src.exists():
                copy_file(src, evidence_dir / name)
                evidence_count += 1

        if origin_count == 0:
            write_text(
                origin_dir / "origin_not_available.md",
                f"# Origin not available\n\n{step.origin_note}\n",
            )
        write_text(
            root / "STEP_README.md",
            f"# {step.title}\n\n"
            f"## 目的\n{step.purpose}\n\n"
            f"## 输入\n{step.inputs}\n\n"
            f"## 输出\n{step.outputs}\n\n"
            f"## Origin 对照说明\n{step.origin_note}\n\n"
            f"## 子目录\n"
            f"- `origin_available/`: 当前能找到的 AZ/origin 文件。\n"
            f"- `our_outputs/`: 最新 workflow run 生成的该阶段结果。\n"
            f"- `comparison_examples/`: origin 和我们的输出按文件名成组放置。\n"
            f"- `evidence/`: 与该阶段相关的审计/检查文件。\n",
        )
        step_counts[step.key] = {
            "origin": origin_count,
            "generated": generated_count,
            "comparison": comp_count,
            "evidence": evidence_count,
            "total": origin_count + generated_count + comp_count + evidence_count + 2,
        }

    copy_file(RUN / "analysis/er_core_workflow.Rmd", OUT_DIR / "00_workflow_source" / "er_core_workflow.Rmd")
    copy_file(RUN / "config/er_workflow_spec.yaml", OUT_DIR / "00_workflow_source" / "er_workflow_spec.yaml")
    copy_file(RUN / "config/study_paths.yaml", OUT_DIR / "00_workflow_source" / "study_paths.yaml")
    for rel in (
        "skills/er-individual-pk-pd-review/code_corpus/az_mock01_core2_reference_plotters.R",
        "skills/er-exposure-response-exploration/code_corpus/az_mock01_core4_er_plotters.R",
        "skills/er-statistical-modeling/code_corpus/az_mock01_core5_km_plotters.R",
    ):
        copy_file(BUNDLE / rel, OUT_DIR / "00_workflow_source" / "az_extracted_plotter_tools" / Path(rel).name)

    final_dir = OUT_DIR / "07_final_human_report"
    ensure_dir(final_dir)
    for src in sorted(HUMAN_REPORT.glob("*.docx")):
        copy_file(src, final_dir / src.name)
    write_text(
        final_dir / "README.md",
        "# Final human report\n\n"
        "这里放的是 workflow 最终给人看的中文 DOCX 报告结果。"
        "它们不是机器审计文件，而是会议/沟通用入口。\n",
    )

    return step_counts


def add_step_sections(doc: Document) -> None:
    doc.add_heading("逐步结果说明", level=1)
    for step in STEPS:
        doc.add_heading(step.title, level=2)
        add_callout(doc, "这一阶段在做什么", step.purpose)
        rows = [
            ("输入", step.inputs),
            ("我们的输出", step.outputs),
            ("origin 情况", step.origin_note),
            ("包内位置", f"steps/{step.key}/"),
        ]
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for left, right in rows:
            cells = table.add_row().cells
            set_cell_text(cells[0], left, True, 8.6)
            set_cell_shading(cells[0], "F2F4F7")
            set_cell_text(cells[1], right, False, 8.6)
        for rel in step.representative_files[:2]:
            add_csv_preview(doc, RUN / rel)


def add_representative_pairs(doc: Document) -> None:
    doc.add_heading("代表性 origin vs 我们输出", level=1)
    pairs = [
        ("Core2 profile preview", "20250925_pkind4"),
        ("Core2 swimmer preview", "swimmer_high_dose"),
        ("Core4 ER pair figure", "ER_AUC1_Res1_efficacy"),
        ("Core5 KM figure", "PFS_KM_by_dose"),
    ]
    for caption, basename in pairs:
        origin = LATEST / f"{basename}__original.png"
        generated = LATEST / f"{basename}__{RUN_LABEL}.png"
        add_image_pair(doc, origin, generated, caption)
    doc.add_heading("代表性表格对照", level=2)
    origin_csv = LATEST / "Final_Logistic_Regression_Complete_Results__original.csv"
    generated_csv = LATEST / f"Final_Logistic_Regression_Complete_Results__{RUN_LABEL}.csv"
    doc.add_paragraph("Origin 表格预览")
    add_csv_preview(doc, origin_csv, max_rows=4, max_cols=5)
    doc.add_paragraph("我们的输出表格预览")
    add_csv_preview(doc, generated_csv, max_rows=4, max_cols=5)


def build_docx(step_counts: dict[str, dict[str, int]]) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc)
    doc.add_heading("Workflow 一共有几步", level=1)
    add_step_overview_table(doc, step_counts)
    add_callout(
        doc,
        "最终 workflow 结果",
        "最终人类可读报告放在 `07_final_human_report/`。其中 `Mock01_项目说明与结果概览.docx` 是主入口；"
        "`主结果_表格对照.docx` 保存结果表 origin/generated 对照；`报告示例_Origin对照.docx` 保存每类图的代表性对照。"
    )
    add_step_sections(doc)
    add_representative_pairs(doc)
    doc.save(DOCX)


def write_readme(step_counts: dict[str, dict[str, int]]) -> None:
    rows = [
        "| Step | Folder | Origin files | Our output files | Comparison/evidence files |",
        "|---|---:|---:|---:|---:|",
    ]
    for step in STEPS:
        counts = step_counts[step.key]
        rows.append(
            f"| {step.title} | `steps/{step.key}` | {counts['origin']} | "
            f"{counts['generated']} | {counts['comparison'] + counts['evidence']} |"
        )
    write_text(
        OUT_DIR / "README_先读这个.md",
        "# Mock01 workflow 分步骤结果对照包\n\n"
        "先打开 `Workflow_分步骤结果对照说明.docx`。如果要看真实文件，按 `steps/01...06` 顺序打开。"
        "每一步下面都有 `origin_available/`、`our_outputs/`、`comparison_examples/`、`evidence/`。\n\n"
        "最终中文报告结果在 `07_final_human_report/`。\n\n"
        + "\n".join(rows)
        + "\n",
    )


def zip_package() -> None:
    if ZIP.exists():
        ZIP.unlink()
    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT_DIR.rglob("*")):
            if path.is_file() and not any(part.startswith("render_") for part in path.parts):
                zf.write(path, path.relative_to(OUT_DIR.parent))


def main() -> None:
    step_counts = build_package_tree()
    build_docx(step_counts)
    write_readme(step_counts)
    zip_package()
    print(OUT_DIR)
    print(ZIP)


if __name__ == "__main__":
    main()
